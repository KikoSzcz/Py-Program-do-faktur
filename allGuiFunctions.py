import allFileFunctions as aff
import tkinter as tk
from tkcalendar import Calendar, DateEntry
from docx import Document
import os
from docx.shared import Pt, Inches
import convertToPolish as conv
from datetime import datetime, timedelta
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path
from tkinter import messagebox

#Zmienne globalne
mainWindow = tk.Tk(className='Program do faktur')
exhibitorVariable = tk.StringVar()
buyerVariable = tk.StringVar()
payOptionVariable = tk.StringVar()
callendar = tk.StringVar()
place = tk.StringVar()
name = tk.StringVar()
path = tk.StringVar()
addedServicesNameArray = []
addedServicesUnitArray = []
addedServicesPriceArray = []
addedServicesQuantityArray = []
addedServicesArray = []


def funCreateLabel(mainWindow):
    #Tworzymy tekst stały aplikacji
    exhibitorLabel = tk.Label(mainWindow, text="Wystawca:")
    #Umieszczamy tekst w konkretnym miejscu w aplikacji
    exhibitorLabel.place(x=10, y=10)

    buyerLabel = tk.Label(mainWindow, text="Nabywca:")
    buyerLabel.place(x=10, y=70)

    payOptionLabel = tk.Label(mainWindow, text="Opcja zapłaty:")
    payOptionLabel.place(x=10, y=130)

    dataLabel = tk.Label(mainWindow, text="Data wystawienia:")
    dataLabel.place(x=10, y=190)

    placeLabel = tk.Label(mainWindow, text="Miejsce wystawienia:")
    placeLabel.place(x=10, y=240)

    nameLabel = tk.Label(mainWindow, text="Numer faktury:")
    nameLabel.place(x=10, y=290)

    pathLabel = tk.Label(mainWindow, text="Ścieżka zapisu:")
    pathLabel.place(x=10, y=420)

    servicesLabel = tk.Label(mainWindow, text="Wybierz usługę:")
    servicesLabel.place(x=400, y=10)

def funCreateOptionMenu(mainWindow):
    #optionMenu exhibitor
    #Pobranie wartości z pliku
    exhibitorArray = aff.funReadFromFile("exhibitor.txt")
    #Stworzenie listy rozwijanej
    exhibitorOptionMenu = tk.OptionMenu(mainWindow, exhibitorVariable, *exhibitorArray)
    #Konfiguracja listy
    exhibitorOptionMenu.config(width=40, anchor='w', height= 1)
    #Umieszczenie listy w oknie apliakcji
    exhibitorOptionMenu.place(x=10, y=30)

    #optionMenu buyer
    buyerArray = []
    buyerArray = aff.funReadFromFile("buyer.txt")
    buyerOptionMenu = tk.OptionMenu(mainWindow, buyerVariable, *buyerArray)
    buyerOptionMenu.config(width=40, anchor='w', height= 1)
    buyerOptionMenu.place(x=10, y=90)

    #OptionMenu pay option
    payOptionArray = []
    payOptionArray = aff.funReadFromFile("payOption.txt")
    payOptionArray.sort()
    payOptionOptionMenu = tk.OptionMenu(mainWindow, payOptionVariable, *payOptionArray)
    payOptionOptionMenu.config(width=40, anchor='w', height= 1)
    payOptionOptionMenu.place(x=10, y=150)

    #Funkcja dodawnia nowych elementów do istniejącej listy i pliku (bazy)
    def funAdd(fileName, data, Window, tkMenu, tkVariable, menuArray):
        data = data[:-1]
        aff.funAddDataToFile(fileName, data)
        menuArray.append(data)
        tkMenu.children["menu"].add_command(label=data, command=tk._setit(tkVariable,data))
        Window.destroy()

    #Funkcja pokazująca okno w którym można wpisać wartość do dodania do listy i pliku (bazy)
    def funAddMenu(text, tkMenu, tkVariable, pathToFile, menuArray):
        #Stworzenie nowego okna
        newWindow = tk.Tk(className='Dodawanie ')
        #Wymiary okna
        newWindow.geometry("665x230")
        #Tekst stały
        funLabel = tk.Label(newWindow, text=text)
        funLabel.place(x=10, y=10)
        #Textbox do którego podaje się nowy element
        funTextBox = tk.Text(newWindow, height=10, width=80)
        funTextBox.place(x=10, y=30)
        #Przyisk wywołujący funAdd w celu dodania wartości z textbox do listy rozwijanej i pliku (bazy)
        funButtonOk = tk.Button(newWindow, text="Dodaj", command=lambda: funAdd(pathToFile, funTextBox.get("1.0", "end"), newWindow, tkMenu, tkVariable, menuArray))
        funButtonOk.place(x=10, y=200)
        newWindow.mainloop()

    #Funkcja usuwania wybranego elementu z listy i pliku (bazy)
    def funDeleteMenu(tkMenu, tkVariable, pathToFile, menuArray):
        #Jeżeli lista ma wybrany jakiś element
        if tkVariable.get()!="":
            #Idzemy po całej liście
            for x in range(len(menuArray)):
                #Aż znajdziemy element do usunięcia
                if menuArray[x] == str(tkVariable.get()):
                    #Zapisujemy jego id (x)
                    r_index = x
                    #Usuwamy element z listy
                    tkMenu["menu"].delete(r_index)
                    #Usuwamy element z pliku
                    aff.funDeleteElementFromFlie(pathToFile, str(tkVariable.get()))
                    #Ustawiamy listę na default
                    tkVariable.set("")


    #Tworzy przyciski do dodwania i usuwania elementów
    def funCreateButtons(mainWindow):
        #Deklaracja przycisku
        exhibitorButton = tk.Button(mainWindow, text="Dodaj", command=lambda: funAddMenu("Wpisz nowego wystawcę:", exhibitorOptionMenu, exhibitorVariable, "exhibitor.txt", exhibitorArray))
        #Umiejscowienie przycisku
        exhibitorButton.place(x=300, y=32)
        exhibitorButton = tk.Button(mainWindow, text="Usuń", command=lambda: funDeleteMenu(exhibitorOptionMenu, exhibitorVariable, "exhibitor.txt", exhibitorArray))
        exhibitorButton.place(x=345, y=32)

        buyerButton = tk.Button(mainWindow, text="Dodaj", command=lambda: funAddMenu("Wpisz nowego nabywcę: ", buyerOptionMenu, buyerVariable, "buyer.txt", buyerArray))
        buyerButton.place(x=300, y=92)
        buyerButton = tk.Button(mainWindow, text="Usuń", command=lambda: funDeleteMenu(buyerOptionMenu, buyerVariable, "buyer.txt", buyerArray))
        buyerButton.place(x=345, y=92)

        payOptionButton = tk.Button(mainWindow, text="Dodaj", command=lambda: funAddMenu("Wpisz nową opcję zapłaty: ", payOptionOptionMenu, payOptionVariable, "payOption.txt", payOptionArray))
        payOptionButton.place(x=300, y=152)
        payOptionButton = tk.Button(mainWindow, text="Usuń", command=lambda: funDeleteMenu(payOptionOptionMenu, payOptionVariable, "payOption.txt", payOptionArray))
        payOptionButton.place(x=345, y=152)

    funCreateButtons(mainWindow)

def funCreateCallendar(mainWindow):
    #Deklaracja kalendarza
    cal = DateEntry(mainWindow,textvariable=callendar, width=12, background='gray', foreground='white', borderwidth=0)
    #Umiejscowienie go
    cal.place(x=13, y=212)

def funCreateEntry(mainWindow):
    #Place
    #Wszytanie danych z pliku
    placeFromFile = aff.funReadLine("place.txt")
    #Deklaracja Entry
    placeEntry = tk.Entry(mainWindow, textvariable = place, width=39, font=('Arial',10))
    #Wpisanie w Entry wartości z zmiennej placeFromFile
    placeEntry.insert(0, placeFromFile)
    #Umiejscowienie elementu w oknie aplikacji
    placeEntry.place(x=13, y=260)

    #FV name
    # Deklaracja Entry
    nameEntry = tk.Entry(mainWindow, textvariable = name, width=39, font=('Arial', 10))
    # Umiejscowienie elementu w oknie aplikacji
    nameEntry.place(x=13, y=310)

    #Path to file entry
    pathFromFile = aff.funReadLine("pathToSave.txt")
    pathEntry = tk.Entry(mainWindow, textvariable = path, width=39, font=('Arial', 10))
    pathEntry.insert(0, pathFromFile)
    pathEntry.place(x=13, y=440)

def funCreateService(mainWindow):
    #services
    servicesArray = []
    servicesArray = aff.funReadServiceFromFile("services.txt")
    servicesVariable = tk.StringVar()
    servicesVariable.set("---")
    addedServicesButtonArray = []

    def funDeleteService(value):
        #obliczenie ID usługi do usunięcia za pomocą zmiennej value (jest to kordynat y)
        toDelete = (value-92)/25
        #Przeniesienie innych elementów od elementu który ma zostać usunięty
        for x in range(int(toDelete), (len(addedServicesArray)-1)):

            addedServicesArray[x] = addedServicesArray[x+1]

            addedServicesNameArray[x].delete(0, len(addedServicesNameArray[x].get()))
            addedServicesNameArray[x].insert(0, addedServicesNameArray[x+1].get())

            addedServicesUnitArray[x].delete(0, len(addedServicesUnitArray[x].get()))
            addedServicesUnitArray[x].insert(0, addedServicesUnitArray[x + 1].get())

            addedServicesPriceArray[x].delete(0, len(addedServicesPriceArray[x].get()))
            addedServicesPriceArray[x].insert(0, addedServicesPriceArray[x + 1].get())

            addedServicesQuantityArray[x].delete(0, len(addedServicesQuantityArray[x].get()))
            addedServicesQuantityArray[x].insert(0, addedServicesQuantityArray[x + 1].get())

            #Obliczenie wartości kordy elementów przeniesionych
            fncy = 95 + (25 * x)
            addedServicesNameArray[x].place(x=402, y=fncy)
            addedServicesUnitArray[x].place(x=642, y=fncy)
            addedServicesPriceArray[x].place(x=752, y=fncy)
            addedServicesQuantityArray[x].place(x=832, y=fncy)
            addedServicesButtonArray[x].place(x=900, y=fncy-3)

        #Usunięcie ostanich elementów z listy
        addedServicesNameArray[-1].destroy()
        addedServicesUnitArray[-1].destroy()
        addedServicesPriceArray[-1].destroy()
        addedServicesQuantityArray[-1].destroy()
        addedServicesButtonArray[-1].destroy()
        del addedServicesArray[-1]
        del addedServicesNameArray[-1]
        del addedServicesUnitArray[-1]
        del addedServicesPriceArray[-1]
        del addedServicesQuantityArray[-1]
        del addedServicesButtonArray[-1]

    def funServiceChangeValue(self):
        #Jeżeli w liscie została zmieniona wartość na inną niż domyślna
        if str(servicesVariable.get()) != "---":
            #Dodanie do tablicy elementu wybranego
            addedServicesArray.append(str(servicesVariable.get()).split("     "))
            #Jeśli jest to pierwszy raz kiedy wartość w liście została zmieniona
            if len(addedServicesArray) == 1:
                #Dodanie tekstu stałego
                servicesTableNameLabe = tk.Label(mainWindow, text="Nazwa usługi")
                servicesTableUnitLabe = tk.Label(mainWindow, text="Jednostka")
                servicesTablePriceLabe = tk.Label(mainWindow, text="Cena")
                servicesTableQuantityLabe = tk.Label(mainWindow, text="Ilość")
                servicesTableNameLabe.place(x=400, y=70)
                servicesTableUnitLabe.place(x=640, y=70)
                servicesTablePriceLabe.place(x=750, y=70)
                servicesTableQuantityLabe.place(x=830, y=70)

            #Z racji na ograniczenie max 15 usług ten if sprawdza czy nie przekroczyliśmy tego zakresu
            if len(addedServicesArray) <= 15:
                #Dodawanie nowych elementów i przypisywanie ich do tablicy obiektów oraz umiejscowienie ich w oknie apliakcji
                fncy = 95 + (25*(len(addedServicesArray)-1))
                addedServicesNameArray.append(tk.Entry(mainWindow, width=37))
                addedServicesNameArray[-1].insert(0, addedServicesArray[-1][1])
                addedServicesNameArray[-1].place(x=402, y=fncy)

                addedServicesUnitArray.append(tk.Entry(mainWindow, width=15))
                addedServicesUnitArray[-1].insert(0, addedServicesArray[-1][2])
                addedServicesUnitArray[-1].place(x=642, y=fncy)

                addedServicesPriceArray.append(tk.Entry(mainWindow, width=10))
                addedServicesPriceArray[-1].place(x=752, y=fncy)

                addedServicesQuantityArray.append(tk.Entry(mainWindow, width=10))
                addedServicesQuantityArray[-1].place(x=832, y=fncy)

                test = fncy-3
                #Dodanie przycisku usuń który po kliknięciu usunie cały wiersz usług
                addedServicesButtonArray.append(tk.Button(mainWindow, text="Usuń", command= lambda: funDeleteService(test)))
                addedServicesButtonArray[-1].config(font=('Arial', 8))
                addedServicesButtonArray[-1].place(x=900, y=(fncy-3))
            #Ustawienie listy na default value
            servicesVariable.set("---")


    #Wywoływanie funkvcji funServiceChangeValue przy zmianie wartości listy
    servicesOptionalMenu = tk.OptionMenu(mainWindow, servicesVariable, *servicesArray, command=funServiceChangeValue)
    servicesOptionalMenu.config(width=40, anchor='w')
    servicesOptionalMenu.place(x=400, y=30)

    def buttonSerivceCombineFunc(Variable, OptionMenu):
        tk._setit(Variable, OptionMenu)
        servicesVariable.set(Variable)
        funServiceChangeValue(Variable)

    def funAddServices(fileName, data1, data2, Window):
        servicesToMenu = str(len(servicesArray)+1) + '     ' + data1 + '     ' + data2
        servicesArray.append(servicesToMenu)
        servicesOptionalMenu.children["menu"].add_command(label=servicesToMenu, command=lambda : buttonSerivceCombineFunc(servicesToMenu, servicesToMenu))
        aff.funAddServiceToFile(fileName, data1, data2)
        Window.destroy()


    def addNewServices():
        newWindow = tk.Tk(className='Dodawanie nowej usługi')
        newWindow.geometry("400x100")
        funLabel = tk.Label(newWindow, text="Wpisz nazwę usługi i jednostkę")
        funLabel.place(x=10, y=10)
        funEntryBox1 = tk.Entry(newWindow, width=40)
        funEntryBox1.place(x=10, y=30)
        funEntryBox2 = tk.Entry(newWindow, width=20)
        funEntryBox2.place(x=260, y=30)
        funButtonOk = tk.Button(newWindow, text="Dodaj", command=lambda: funAddServices("services.txt", funEntryBox1.get(), funEntryBox2.get(), newWindow))
        funButtonOk.place(x=10, y=55)
        newWindow.mainloop()

    def funDeleteSelectedService(data, window):
        if data != '---':
            funTab = data.split('     ')
            servicesArray.remove(data)
            servicesOptionalMenu["menu"].delete(int(funTab[0])-1)
            aff.funDeleteServiceFromFile("services.txt", funTab[1], funTab[2])
            window.destroy()


    def deleteSelectedService():
        newWindow = tk.Tk(className="Usuwanie usługi")
        newWindow.geometry("300x100")
        funLabel = tk.Label(newWindow, text="Wybierz usługę do usunięcia")
        funLabel.place(x=10, y=10)
        funDeleteVariable = tk.StringVar()
        funDeleteVariable.set("---")
        funArray = servicesArray
        funOptionMenu = tk.OptionMenu(newWindow, funDeleteVariable, *funArray)
        funOptionMenu.config(width=40 , anchor='w', height= 1)
        funOptionMenu.place(x=10, y=30)

        funButton = tk.Button(newWindow, text="Usuń wybrane", command=lambda: funDeleteSelectedService(funDeleteVariable.get(), newWindow))
        funButton.place(x=12, y=62)
        newWindow.mainloop()

    #Services button
    servicesAddButton = tk.Button(mainWindow, text="Dodaj", command=lambda : addNewServices())
    servicesAddButton.place(x=690, y=32)
    servicesDeleteButton = tk.Button(mainWindow, text="Usuń", command=lambda : deleteSelectedService())
    servicesDeleteButton.place(x=735, y=32)

def funCreateDocument(funString):
    exhibitor = exhibitorVariable.get()
    buyer = buyerVariable.get()
    payOption = payOptionVariable.get()
    date = str(callendar.get()).split('/')
    datetime_object = datetime.strptime(str(callendar.get()), '%m/%d/%y')
    city = place.get()
    vfname=name.get()
    pathToFile=path.get()

    dateDay = date[1]
    dateMonth = date[0]
    dateYear = "20"+date[2]
    if int(date[1])<10:
        dateDay ="0"+date[1]
    if int(date[1])<10:
        dateMonth ="0"+date[0]

    oldDateDay = dateDay
    oldDateMonth = dateMonth
    oldDateYear = dateYear

    #Tablica usług
    servicesArray = []
    for x in range(len(addedServicesNameArray)):
        krotka = (addedServicesNameArray[x].get(), addedServicesUnitArray[x].get(), addedServicesPriceArray[x].get(), addedServicesQuantityArray[x].get())
        servicesArray.append(krotka)

    #servicesArray[x][0] - nazwa
    #servicesArray[x][1] - jednostka
    #servicesArray[x][2] - cena
    #servicesArray[x][3] - ilość


    document = Document()
    firstLine = "WYSTAWCA"
    firstLineBreak=""

    #Pierwsza linia
    p = document.add_paragraph()
    runner = p.add_run(firstLine)
    runner.bold = True
    runner.underline = True
    runner.font.name = "Times News Roman"
    runner.font.size = Pt(12)
    for x in range(81-len(city)):
        firstLineBreak+=" "
    runner = p.add_run(firstLineBreak+city.upper()+" "+dateDay+"."+dateMonth+"."+dateYear+"r.\n")
    runner.font.name = "Times News Roman"
    runner.font.size = Pt(12)
    runner.bold = True

    #Wystawca
    exhib = exhibitor.split('\n')
    exhibOther = ""
    for x in range(1,len(exhib)):
        exhibOther += exhib[x]+"\n"
    runner = p.add_run(exhib[0]+"\n")
    runner.font.name = "Arial Black"
    runner.font.size = Pt(12)
    runner.bold = True
    runner = p.add_run(exhibOther)
    runner.font.name = "Times News Roman"
    runner.font.size = Pt(12)
    runner.bold = True

    #Nabywca
    runner = p.add_run("\nNABYWCA\n")
    runner.font.name = "Times News Roman"
    runner.font.size = Pt(12)
    runner.bold = True
    runner.underline = True
    runner = p.add_run(buyer)
    runner.font.name = "Times News Roman"
    runner.font.size = Pt(12)
    runner.bold = True

    #Numer faktury
    fv = document.add_paragraph()
    fv.alignment = 1
    runnerFv = fv.add_run("\n\nRACHUNEK NR "+vfname+" - " + funString.upper())
    runnerFv.font.name = "Times News Roman"
    runnerFv.font.size = Pt(12)
    runnerFv.bold = True

    #Tabela
    shading_elm01 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm02 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm03 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm04 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm05 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm06 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading_elm07 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm01)
    hdr_cells[0].text = 'L.P.'
    hdr_cells[0].width = Inches(0.85/2.54)


    hdr_cells[1].text = 'NAZWA USŁUGI'
    hdr_cells[1].width = Inches(6.57/2.54)
    hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm02)

    hdr_cells[2].text = 'J.M.'
    hdr_cells[2].width = Inches(1.51/2.54)
    hdr_cells[2]._tc.get_or_add_tcPr().append(shading_elm03)

    hdr_cells[3].text = 'ILOŚĆ'
    hdr_cells[3].width = Inches(1.46/ 2.54)
    hdr_cells[3]._tc.get_or_add_tcPr().append(shading_elm04)

    hdr_cells[4].text = 'CENA JEDNOSTKOWA'
    hdr_cells[4].width = Inches(2.93/2.54)
    hdr_cells[4]._tc.get_or_add_tcPr().append(shading_elm05)

    hdr_cells[5].text = 'WARTOŚĆ'
    hdr_cells[5].width = Inches(2.7/2.54)
    hdr_cells[5]._tc.get_or_add_tcPr().append(shading_elm06)

    for x in range(0,6):
        hdr_cells[x].paragraphs[0].alignment = 1
        paragraph = hdr_cells[x].paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(10)
        font.name = "Times News Roman"

    sum = 0.0

    for x in range(len(servicesArray)):
        row_cells = table.add_row().cells
        row_cells[0].text=str(x+1)
        row_cells[1].text=str(servicesArray[x][0])
        row_cells[2].text=str(servicesArray[x][1])
        row_cells[3].text=str('%.2f' % float(servicesArray[x][3]))
        row_cells[4].text=str('%.2f' % float(servicesArray[x][2]))
        row_cells[5].text=str('%.2f' % (float(servicesArray[x][3])*float(servicesArray[x][2])))
        sum += float(servicesArray[x][3])*float(servicesArray[x][2])
        row_cells[0].width = Inches(0.85/2.54)
        row_cells[1].width = Inches(6.57/2.54)
        row_cells[2].width = Inches(1.51/2.54)
        row_cells[3].width = Inches(1.46/ 2.54)
        row_cells[4].width = Inches(2.93/2.54)
        row_cells[5].width = Inches(2.7/2.54)

        for y in range(0,6):
            row_cells[y].paragraphs[0].alignment = 1
            paragraph = row_cells[y].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(12)
            font.name = "Times News Roman"

        row_cells[5].paragraphs[0].alignment = 2

    temp = document.add_paragraph()
    temp.line_spacing = Pt(0)
    runnerTemp = temp.add_run("")
    runnerTemp.font.name = "Times News Roman"
    runnerTemp.font.size = Pt(1)

    #Tabelka sumy
    sumtable = document.add_table(rows=1, cols=2)
    sumtable.style = 'Table Grid'
    sumtable.alignment = 2
    hdr_cells = sumtable.rows[0].cells

    hdr_cells[0].text = 'RAZEM'
    hdr_cells[0].width = Inches(2.92 / 2.54)

    hdr_cells[1].text = str('%.2f' % sum)+" PLN"
    hdr_cells[1].width = Inches(4.16 / 2.54)

    for x in range(0, 2):
        hdr_cells[x].paragraphs[0].alignment = 1
        paragraph = hdr_cells[x].paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size = Pt(12)
        font.name = "Times News Roman"
        if x == 0:
            font.bold = True
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm07)
    hdr_cells[1].paragraphs[0].alignment = 2


    #Inne wartości
    sumStr = str('%.2f' % sum)
    sumInt = sumStr.split('.')
    inPolish = conv.słownie(int(sumInt[0]), skala='krótka')



    other = document.add_paragraph()
    runnerOther = other.add_run("\nDO ZAPŁATY: " + str('%.2f' % sum) + " PLN\nDO ZAPŁATY SŁOWNIE: "+inPolish.upper()+" PLN "+sumInt[1]+"/100\n")
    runnerOther.font.name = "Times News Roman"
    runnerOther.font.size = Pt(12)
    runnerOther.bold = True
    if(payOption=='gotówka'):
        runnerOther = other.add_run("SPOSÓB ZAPŁATY: "+payOption.upper()+"\nTERMIN ZAPŁATY: "+" "+dateDay+"."+dateMonth+"."+dateYear+"r.\n")
        runnerOther.font.name = "Times News Roman"
        runnerOther.font.size = Pt(12)
        runnerOther.bold = True
    else:
        payOption = payOption.split(' ')
        datetime_object = datetime_object + timedelta(days=int(payOption[1]))
        dateDay = str(datetime_object.day)
        dateMonth = str(datetime_object.month)
        dateYear = str(datetime_object.year)
        if int(dateDay)<10:
            dateDay = "0"+dateDay
        if int(dateMonth)<10:
            dateMonth = "0"+dateMonth
        runnerOther = other.add_run("SPOSÓB ZAPŁATY: "+payOption[0].upper()+" - "+ payOption[1]+" DNI\nTERMIN ZAPŁATY: " + dateDay + "." + dateMonth + "." + dateYear + "r.\n")
        runnerOther.font.name = "Times News Roman"
        runnerOther.font.size = Pt(12)
        runnerOther.bold = True

    runnerOther = other.add_run("\n\nWSTAWACA ZWOLNIONY PODMIOTOWO Z PODATKU VAT (PODATKU OD TOWARÓW I USŁUG)")
    runnerOther.font.name = "Times News Roman"
    runnerOther.font.size = Pt(9)
    runnerOther.bold = False

    signature = document.add_paragraph()
    runnerSignature = signature.add_run("\n\n\n\n\n\nPIECZĘĆ I PODPIS WYSTAWCY\n")
    runnerSignature.font.name = "Times News Roman"
    runnerSignature.font.size = Pt(9)
    runnerSignature.bold = False
    signature.alignment = 2

    bottomText = document.add_paragraph()
    runnerbottomText = bottomText.add_run("ŚWIADCZYMY USŁUGI: KOPARKAMI, ŁADOWARKAMI, ZAGĘSZCZARKIAMI  I WIELE INNYCH. WIĘCEJ INFORMACJI NA STRONIE www.valdi1999.pl")
    runnerbottomText.font.name = "Times News Roman"
    runnerbottomText.font.size = Pt(6)
    runnerbottomText.bold = False
    bottomText.alignment = 1

    if funString == "oryginał":
        nameSaved = oldDateYear + "-" + oldDateMonth + "-" + oldDateDay + "-" + vfname + ".docx"
    else:
        nameSaved = oldDateYear + "-" + oldDateMonth + "-" + oldDateDay + "-" + vfname + "-kopia.docx"
    document.save(pathToFile+"\\"+nameSaved)
    os.startfile(pathToFile+"\\"+nameSaved)

def funEditPlacePath(fileName):
    #Otwarcie pliku przez notatnik i czekanie aż notatnik zostanie zamknięty
    os.system(fileName)
    #Wczytanie nowej wartości z dokumentu teksotwego
    if fileName == "place.txt":
        place.set(str(aff.funReadLine(fileName)))
    if fileName == "pathToSave.txt":
        path.set(str(aff.funReadLine(fileName)))

def funCreatePlacePathButton():
    #Stworzenie przycisku do edycji domyślnej wartości pola miejscowość
    placeButton = tk.Button(mainWindow, text="Edytuj", command=lambda: funEditPlacePath("place.txt"))
    placeButton.place(x=299, y=257)

    # Stworzenie przycisku do edycji domyślnej wartości pola ścieżka zapisu
    pathButton = tk.Button(mainWindow, text="Edytuj", command=lambda: funEditPlacePath("pathToSave.txt"))
    pathButton.place(x=299, y=437)

def funStartToCreateDocx():
    date = str(callendar.get()).split('/')
    datetime_object = datetime.strptime(str(callendar.get()), '%m/%d/%y')
    pathToFile = path.get()
    vfname = name.get()

    dateDay = date[1]
    dateMonth = date[0]
    dateYear = "20" + date[2]
    if int(date[1]) < 10:
        dateDay = "0" + date[1]
    if int(date[1]) < 10:
        dateMonth = "0" + date[0]
    nameSaved = dateYear + "-" + dateMonth + "-" + dateDay + "-" + vfname + ".docx"

    my_file = Path(pathToFile+"\\"+nameSaved)
    if my_file.is_file():
        MsgBox = tk.messagebox.askquestion('Istnieje taki dokument', 'Istnieje już taki dokument o nazwie '+nameSaved+', czy chcesz go nadpisać?')
        if MsgBox == 'yes':
            funCreateDocument("oryginał")
            funCreateDocument("kopia")
    else:
        funCreateDocument("oryginał")
        funCreateDocument("kopia")

def funStart():
    # Rozmiar okna
    mainWindow.geometry("950x480")

    funCreateLabel(mainWindow)
    funCreateOptionMenu(mainWindow)
    funCreateCallendar(mainWindow)
    funCreateEntry(mainWindow)
    funCreateService(mainWindow)
    funCreatePlacePathButton()
    generateDocumentBtn = tk.Button(mainWindow, text="Generuj dokument tekstowy", command=lambda : funStartToCreateDocx())
    generateDocumentBtn.config(font=('Arial', 11))
    generateDocumentBtn.place(x=50, y=365)
    # Wywołanie okna
    mainWindow.mainloop()